from docx import Document
from docx.shared import Inches

# Create a new Word document
doc = Document()

# Add user's photo at the top
doc.add_picture('./profile picture.jpeg', width=Inches(1.5))

# Add name and contact information
doc.add_heading('Rodolfo Ravelli', level=1)
doc.add_paragraph('Email: your.email@example.com | Phone: (XX) XXXXX-XXXX | LinkedIn: linkedin.com/in/yourprofile')

# Add sections
doc.add_heading('Professional Summary', level=2)
doc.add_paragraph('Experienced professional with a strong background in [Your Field]. Skilled in [Key Skills]. Seeking to leverage expertise in [Area] to contribute to [Company Name].')

doc.add_heading('Work Experience', level=2)
experience = doc.add_paragraph()
experience.add_run('Job Title at Company Name').bold = True
experience.add_run(' | Location | Month Year – Present\n')
experience.add_run('- Describe your responsibilities and achievements here, using bullet points or short sentences.\n- Another key responsibility or accomplishment.')

doc.add_heading('Education', level=2)
education = doc.add_paragraph()
education.add_run('Degree, Major').bold = True
education.add_run(' | University Name, Location | Month Year – Month Year\n')
education.add_run('- Any relevant coursework or honors.')

doc.add_heading('Skills', level=2)
skill_list = doc.add_paragraph()
skill_list.add_run('- Skill 1\n- Skill 2\n- Skill 3\n- Skill 4')

doc.add_heading('Certifications', level=2)
certifications = doc.add_paragraph()
certifications.add_run('- Certification Name, Issuing Organization, Year\n- Another Certification')

# Save the document
file_path = './Rodolfo_Ravelli_Resume_Template.docx'
doc.save(file_path)

file_path

